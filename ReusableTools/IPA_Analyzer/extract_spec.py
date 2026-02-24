#!/usr/bin/env python3
"""
Functional Specification Data Extraction Tool

Purpose: Extract structured data from ANA-050 documents for Kiro to analyze
This tool does NOT analyze - it only extracts and organizes data

Output: JSON file with requirements, text sections, metadata
"""

import re
import json
from pathlib import Path
try:
    import docx
except ImportError:
    print("Warning: python-docx not installed. Install with: pip install python-docx")
    docx = None


def extract_spec(spec_file: str, output_file: str = None) -> dict:
    """
    Extract structured data from functional specification
    
    Args:
        spec_file: Path to ANA-050 DOCX file
        output_file: Optional output JSON file path
    
    Returns:
        Dictionary with extracted data
    """
    if not docx:
        return {'error': 'python-docx not installed'}
    
    print(f"Extracting data from {Path(spec_file).name}...")
    
    try:
        doc = docx.Document(spec_file)
    except Exception as e:
        return {'error': f'Could not read document: {str(e)}'}
    
    result = {
        'file': Path(spec_file).name,
        'metadata': extract_metadata(doc, spec_file),
        'requirements': extract_requirements(doc),
        'sections': extract_sections(doc),
        'tables': extract_tables(doc),
        'full_text': extract_full_text(doc)
    }
    
    # Save to JSON if output file specified
    if output_file:
        with open(output_file, 'w', encoding='utf-8', errors='replace') as f:
            json.dump(result, f, indent=2)
        print(f"✓ Data saved to {output_file}")
    
    return result


def extract_metadata(doc, file_path: str) -> dict:
    """Extract document metadata"""
    metadata = {
        'filename': Path(file_path).name,
        'paragraph_count': len(doc.paragraphs),
        'table_count': len(doc.tables)
    }
    
    # Try to extract RICE ID from filename
    filename = Path(file_path).name.upper()
    if 'REPORT' in filename:
        metadata['rice_type_hint'] = 'Report'
    elif 'INTERFACE' in filename or 'INTFC' in filename:
        metadata['rice_type_hint'] = 'Interface'
    elif 'CONVERSION' in filename or 'CONV' in filename:
        metadata['rice_type_hint'] = 'Conversion'
    elif 'ENHANCEMENT' in filename or 'ENHANC' in filename:
        metadata['rice_type_hint'] = 'Enhancement'
    else:
        metadata['rice_type_hint'] = 'Unknown'
    
    return metadata


def extract_requirements(doc) -> list:
    """Extract numbered requirements and bullet points"""
    requirements = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip empty or very short paragraphs
        if len(text) < 20:
            continue
        
        # Check if it's a numbered requirement or bullet point
        if re.match(r'^\d+\.', text) or text.startswith('•') or text.startswith('-'):
            # Clean up the text
            cleaned = re.sub(r'^\d+\.\s*', '', text)
            cleaned = re.sub(r'^[•\-]\s*', '', cleaned)
            
            if len(cleaned) > 20 and len(cleaned) < 1000:
                requirements.append({
                    'text': cleaned,
                    'original': text[:100]  # First 100 chars for reference
                })
    
    return requirements[:100]  # Limit to 100 requirements


def extract_sections(doc) -> dict:
    """Extract major sections from document"""
    sections = {}
    current_section = 'Introduction'
    current_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Check if this is a heading (simple heuristic)
        if para.style.name.startswith('Heading') or (len(text) < 100 and text.isupper()):
            # Save previous section
            if current_content:
                sections[current_section] = '\n'.join(current_content)
            
            # Start new section
            current_section = text
            current_content = []
        else:
            if text:
                current_content.append(text)
    
    # Save last section
    if current_content:
        sections[current_section] = '\n'.join(current_content)
    
    return sections


def extract_tables(doc) -> list:
    """Extract tables from document"""
    tables_data = []
    
    for idx, table in enumerate(doc.tables):
        table_data = {
            'table_number': idx + 1,
            'rows': []
        }
        
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data['rows'].append(row_data)
        
        tables_data.append(table_data)
    
    return tables_data


def extract_full_text(doc) -> str:
    """Extract all text from document"""
    return '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python extract_spec.py <spec_file.docx>")
        sys.exit(1)
    
    spec_file = sys.argv[1]
    output_file = 'Temp/spec_data.json'
    
    result = extract_spec(spec_file, output_file)
    
    if 'error' in result:
        print(f"Error: {result['error']}")
    else:
        print(f"\nExtracted:")
        print(f"  - Paragraphs: {result['metadata']['paragraph_count']}")
        print(f"  - Tables: {result['metadata']['table_count']}")
        print(f"  - Requirements: {len(result['requirements'])}")
        print(f"  - Sections: {len(result['sections'])}")
        print(f"  - RICE Type Hint: {result['metadata']['rice_type_hint']}")
